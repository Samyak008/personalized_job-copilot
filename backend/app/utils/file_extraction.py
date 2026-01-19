"""
Resume File Extraction Utilities

Extracts text content from PDF and Word documents for resume parsing.
"""

import io
from pathlib import Path
from typing import Union
from loguru import logger

# PDF extraction
import fitz  # PyMuPDF

# Word document extraction
from docx import Document


class UnsupportedFileTypeError(Exception):
    """Raised when the file type is not supported."""
    pass


class FileExtractionError(Exception):
    """Raised when text extraction fails."""
    pass


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file content.
    
    Args:
        file_content: Raw bytes of the PDF file
        
    Returns:
        Extracted text as string
        
    Raises:
        FileExtractionError: If extraction fails
    """
    try:
        # Open PDF from bytes
        doc = fitz.open(stream=file_content, filetype="pdf")
        
        text_parts = []
        for page_num, page in enumerate(doc):
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(page_text)
            logger.debug(f"Extracted {len(page_text)} chars from PDF page {page_num + 1}")
        
        num_pages = len(doc)
        doc.close()
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from PDF ({num_pages} pages)")
        
        return full_text
        
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {e}")
        raise FileExtractionError(f"PDF extraction failed: {e}") from e


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from Word document (.docx) content.
    
    Args:
        file_content: Raw bytes of the Word document
        
    Returns:
        Extracted text as string
        
    Raises:
        FileExtractionError: If extraction fails
    """
    try:
        # Open document from bytes
        doc = Document(io.BytesIO(file_content))
        
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        full_text = "\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from Word document")
        
        return full_text
        
    except Exception as e:
        logger.error(f"Failed to extract text from Word document: {e}")
        raise FileExtractionError(f"Word document extraction failed: {e}") from e


def extract_text_from_file(
    file_content: bytes, 
    filename: str
) -> str:
    """
    Extract text from a resume file (PDF or Word).
    
    Args:
        file_content: Raw bytes of the file
        filename: Original filename (used to determine file type)
        
    Returns:
        Extracted text as string
        
    Raises:
        UnsupportedFileTypeError: If file type is not supported
        FileExtractionError: If extraction fails
    """
    # Determine file type from extension
    extension = Path(filename).suffix.lower()
    
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Unsupported file type: {extension}. Supported types: {SUPPORTED_EXTENSIONS}"
        )
    
    logger.info(f"Extracting text from {filename} ({extension})")
    
    if extension == ".pdf":
        return extract_text_from_pdf(file_content)
    elif extension in {".docx", ".doc"}:
        # Note: .doc files may not work perfectly with python-docx
        if extension == ".doc":
            logger.warning(".doc format has limited support. Consider converting to .docx")
        return extract_text_from_docx(file_content)
    else:
        raise UnsupportedFileTypeError(f"Unsupported file type: {extension}")


async def extract_text_from_upload(
    file_content: bytes,
    filename: str,
    content_type: str | None = None
) -> str:
    """
    Async wrapper for file extraction (for use in FastAPI endpoints).
    
    Args:
        file_content: Raw bytes of the uploaded file
        filename: Original filename
        content_type: MIME type (optional, for validation)
        
    Returns:
        Extracted text as string
    """
    # Validate content type if provided
    valid_types = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }
    
    if content_type and content_type not in valid_types:
        logger.warning(f"Unexpected content type: {content_type}, proceeding based on filename")
    
    return extract_text_from_file(file_content, filename)
